#!/usr/bin/env python3
"""
Test script for CommentGenerator to verify it works without AttributeError
"""

import os
import sys
from docx import Document
from src.comment_generator import CommentGenerator

def test_comment_generator():
    """Test the CommentGenerator with a sample document"""
    
    # Check if we have a sample document to test with
    sample_doc_path = "examples/sample_documents/sample_contract_compliant.docx"
    
    if not os.path.exists(sample_doc_path):
        print(f"Sample document not found at {sample_doc_path}")
        print("Creating a simple test document...")
        
        # Create a simple test document
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph for the CommentGenerator.')
        doc.add_paragraph('This document will be used to test comment insertion.')
        
        # Ensure directory exists
        os.makedirs(os.path.dirname(sample_doc_path), exist_ok=True)
        doc.save(sample_doc_path)
        print(f"Created test document at {sample_doc_path}")
    
    # Test the CommentGenerator
    try:
        print("Testing CommentGenerator...")
        
        # Create sample issues
        test_issues = [
            {
                'type': 'Jurisdiction Issue',
                'severity': 'High',
                'description': 'Document references incorrect jurisdiction',
                'section': 'Governing Law',
                'adgm_reference': 'ADGM Commercial Regulations 2020, Section 15.2',
                'suggestion': 'Change jurisdiction to Abu Dhabi Global Market'
            },
            {
                'type': 'Missing Clause',
                'severity': 'Medium',
                'description': 'Document missing required compliance clause',
                'section': 'Compliance',
                'adgm_reference': 'ADGM Companies Regulations 2020, Art. 22',
                'suggestion': 'Add compliance and regulatory clause'
            }
        ]
        
        # Initialize CommentGenerator
        comment_gen = CommentGenerator()
        
        # Test adding comments to document
        output_path = "test_output_commented.docx"
        
        print(f"Processing document: {sample_doc_path}")
        print(f"Output will be saved to: {output_path}")
        
        success = comment_gen.add_comments_to_document(
            sample_doc_path, 
            test_issues, 
            output_path
        )
        
        if success:
            print("✅ SUCCESS: CommentGenerator processed document without errors!")
            print(f"✅ Output file created: {output_path}")
            
            # Verify output file exists and has content
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                print(f"✅ Output file size: {file_size} bytes")
                
                # Clean up test output
                os.remove(output_path)
                print("✅ Test output file cleaned up")
            else:
                print("❌ Output file was not created")
                return False
        else:
            print("❌ FAILED: CommentGenerator failed to process document")
            return False
            
        # Test individual methods
        print("\nTesting individual methods...")
        
        # Test with a new document
        test_doc = Document()
        test_doc.add_paragraph("Test paragraph for comment testing.")
        
        # Test adding comment to paragraph
        test_para = test_doc.paragraphs[0]
        comment_success = comment_gen.add_comment_to_paragraph(
            test_para, 
            test_doc, 
            "Test comment", 
            "Test Author"
        )
        
        if comment_success:
            print("✅ SUCCESS: add_comment_to_paragraph works correctly")
        else:
            print("❌ FAILED: add_comment_to_paragraph failed")
            return False
        
        # Test summary comment
        compliance_result = {
            'compliance_score': 75,
            'total_issues': 2,
            'is_compliant': False,
            'high_severity_issues': 1,
            'medium_severity_issues': 1,
            'low_severity_issues': 0
        }
        
        summary_success = comment_gen.generate_summary_comment(test_doc, compliance_result)
        
        if summary_success:
            print("✅ SUCCESS: generate_summary_comment works correctly")
        else:
            print("❌ FAILED: generate_summary_comment failed")
            return False
        
        print("\n🎉 ALL TESTS PASSED! CommentGenerator is working correctly.")
        return True
        
    except Exception as e:
        print(f"❌ ERROR: Test failed with exception: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = test_comment_generator()
    sys.exit(0 if success else 1)
