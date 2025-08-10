#!/usr/bin/env python3
"""
Script to create a sample .docx document for testing the ADGM Corporate Agent.
This creates a sample contract that will be used to test the compliance checker.
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def create_sample_contract():
    """Create a sample contract document for testing."""
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('SERVICE AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_para = doc.add_paragraph()
    date_para.add_run('Date: January 15, 2025').bold = True
    
    # Add parties section
    doc.add_heading('PARTIES', level=1)
    
    party_a = doc.add_paragraph()
    party_a.add_run('Party A: ').bold = True
    party_a.add_run('ABC Corporation Ltd.\n')
    party_a.add_run('Address: Business Bay, Dubai, UAE\n')
    party_a.add_run('Registration Number: ABC123456')
    
    party_b = doc.add_paragraph()
    party_b.add_run('Party B: ').bold = True
    party_b.add_run('XYZ Services LLC\n')
    party_b.add_run('Address: Al Maryah Island, Abu Dhabi, UAE\n')
    party_b.add_run('Registration Number: XYZ789012')
    
    # Add definitions
    doc.add_heading('DEFINITIONS', level=1)
    
    definitions = [
        ('1.1', '"Agreement" means this Service Agreement'),
        ('1.2', '"Effective Date" means the date of execution of this agreement'),
        ('1.3', '"Services" means the consulting services described in Schedule A'),
        ('1.4', '"Term" means the duration of this agreement')
    ]
    
    for num, definition in definitions:
        para = doc.add_paragraph()
        para.add_run(f'{num} ').bold = True
        para.add_run(definition)
    
    # Add scope of services
    doc.add_heading('SCOPE OF SERVICES', level=1)
    
    scope_para = doc.add_paragraph()
    scope_para.add_run('Party B shall provide consulting services to Party A in accordance with the specifications outlined in Schedule A attached hereto.')
    
    # Add term
    doc.add_heading('TERM', level=1)
    
    term_para = doc.add_paragraph()
    term_para.add_run('This agreement shall commence on the Effective Date and continue for a period of twelve (12) months, unless earlier terminated in accordance with the provisions herein.')
    
    # Add payment terms
    doc.add_heading('PAYMENT TERMS', level=1)
    
    payment_terms = [
        'Party A shall pay Party B a monthly fee of AED 50,000 for the services rendered.',
        'Payment shall be made within 30 days of receipt of invoice.',
        'All amounts are exclusive of VAT unless otherwise stated.',
        'Late payments shall incur interest at 5% per annum.'
    ]
    
    for term in payment_terms:
        doc.add_paragraph(term, style='List Bullet')
    
    # Add governing law - THIS IS INCORRECT FOR ADGM COMPLIANCE
    doc.add_heading('GOVERNING LAW AND JURISDICTION', level=1)
    
    # This clause is intentionally non-compliant to test red flag detection
    governing_law = doc.add_paragraph()
    governing_law.add_run('This agreement shall be governed by and construed in accordance with the laws of England and Wales.')
    
    jurisdiction = doc.add_paragraph()
    jurisdiction.add_run('The parties hereby submit to the exclusive jurisdiction of the courts of England and Wales.')
    
    # Add termination
    doc.add_heading('TERMINATION', level=1)
    
    termination_para = doc.add_paragraph()
    termination_para.add_run('Either party may terminate this agreement with thirty (30) days written notice to the other party.')
    
    # Add signatures
    doc.add_heading('SIGNATURES', level=1)
    
    # Party A signature
    doc.add_paragraph('Party A:')
    doc.add_paragraph('Name: _________________________')
    doc.add_paragraph('Title: _________________________')
    doc.add_paragraph('Date: _________________________')
    doc.add_paragraph('Signature: _____________________')
    
    # Party B signature
    doc.add_paragraph('Party B:')
    doc.add_paragraph('Name: _________________________')
    doc.add_paragraph('Title: _________________________')
    doc.add_paragraph('Date: _________________________')
    doc.add_paragraph('Signature: _____________________')
    
    return doc

def create_compliant_contract():
    """Create an ADGM-compliant version of the contract."""
    
    doc = create_sample_contract()
    
    # Replace the non-compliant governing law section
    for paragraph in doc.paragraphs:
        if 'England and Wales' in paragraph.text:
            # Replace with ADGM-compliant text
            if 'governed by and construed' in paragraph.text:
                paragraph.text = 'This agreement shall be governed by and construed in accordance with the laws of the Abu Dhabi Global Market.'
            elif 'exclusive jurisdiction' in paragraph.text:
                paragraph.text = 'The parties hereby submit to the exclusive jurisdiction of the courts of the Abu Dhabi Global Market.'
    
    return doc

if __name__ == "__main__":
    # Create non-compliant version (for testing red flag detection)
    non_compliant_doc = create_sample_contract()
    non_compliant_doc.save('examples/sample_documents/sample_contract_non_compliant.docx')
    print("Created non-compliant sample contract: sample_contract_non_compliant.docx")
    
    # Create compliant version (for testing compliance checking)
    compliant_doc = create_compliant_contract()
    compliant_doc.save('examples/sample_documents/sample_contract_compliant.docx')
    print("Created compliant sample contract: sample_contract_compliant.docx")
    
    print("\nSample documents created successfully!")
    print("These documents can be used to test the ADGM Corporate Agent.")
