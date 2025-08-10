"""
Document Processor Module
Handles parsing, analysis, and processing of .docx documents
"""

import os
from typing import Dict, List, Any, Optional
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import json


class DocumentProcessor:
    """Processes .docx documents for legal analysis"""
    
    def __init__(self):
        self.supported_extensions = ['.docx']
        
    def can_process(self, file_path: str) -> bool:
        """Check if the file can be processed"""
        _, ext = os.path.splitext(file_path)
        return ext.lower() in self.supported_extensions
    
    def extract_text(self, file_path: str) -> str:
        """Extract text content from .docx file"""
        try:
            doc = Document(file_path)
            text_content = []
            
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # Paragraph
                    paragraph = Paragraph(element, doc)
                    if paragraph.text.strip():
                        text_content.append(paragraph.text.strip())
                elif isinstance(element, CT_Tbl):
                    # Table
                    table = Table(element, doc)
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")
    
    def extract_structure(self, file_path: str) -> Dict[str, Any]:
        """Extract document structure including headings, sections, etc."""
        try:
            doc = Document(file_path)
            structure = {
                'title': '',
                'sections': [],
                'tables': [],
                'paragraphs': []
            }
            
            # Extract title (first paragraph)
            if doc.paragraphs:
                structure['title'] = doc.paragraphs[0].text.strip()
            
            # Extract sections and content
            current_section = None
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # Check if this is a heading (simple heuristic)
                if (paragraph.style.name.startswith('Heading') or 
                    len(text) < 100 and text.isupper() or
                    text.endswith(':') and len(text) < 50):
                    current_section = {
                        'title': text,
                        'content': []
                    }
                    structure['sections'].append(current_section)
                elif current_section:
                    current_section['content'].append(text)
                else:
                    structure['paragraphs'].append(text)
            
            # Extract table information
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                structure['tables'].append(table_data)
            
            return structure
        except Exception as e:
            raise Exception(f"Error extracting document structure: {str(e)}")
    
    def identify_document_type(self, content: str, structure: Dict[str, Any]) -> str:
        """Identify the type of legal document based on content and structure"""
        content_lower = content.lower()
        
        # Company Formation Documents
        if any(keyword in content_lower for keyword in ['articles of association', 'memorandum of association']):
            return 'Articles of Association'
        elif 'memorandum' in content_lower and 'association' in content_lower:
            return 'Memorandum of Association'
        elif 'board resolution' in content_lower:
            return 'Board Resolution'
        elif 'shareholder resolution' in content_lower:
            return 'Shareholder Resolution'
        elif 'incorporation application' in content_lower:
            return 'Incorporation Application'
        elif 'ubo declaration' in content_lower:
            return 'UBO Declaration'
        elif 'register of members' in content_lower:
            return 'Register of Members and Directors'
        elif 'change of registered address' in content_lower:
            return 'Change of Registered Address Notice'
        
        # Regulatory and Compliance
        elif any(keyword in content_lower for keyword in ['license', 'licensing', 'regulatory']):
            return 'Regulatory Filing'
        elif any(keyword in content_lower for keyword in ['employment', 'hr', 'contract']):
            return 'Employment Contract'
        elif any(keyword in content_lower for keyword in ['commercial', 'agreement', 'partnership']):
            return 'Commercial Agreement'
        elif any(keyword in content_lower for keyword in ['compliance', 'risk', 'policy']):
            return 'Compliance Policy'
        
        return 'Unknown Document Type'
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Main method to process a document and return comprehensive analysis"""
        if not self.can_process(file_path):
            raise ValueError(f"Unsupported file type. Supported: {self.supported_extensions}")
        
        # Extract content and structure
        text_content = self.extract_text(file_path)
        structure = self.extract_structure(file_path)
        
        # Identify document type
        doc_type = self.identify_document_type(text_content, structure)
        
        return {
            'file_path': file_path,
            'file_name': os.path.basename(file_path),
            'document_type': doc_type,
            'text_content': text_content,
            'structure': structure,
            'word_count': len(text_content.split()),
            'paragraph_count': len(structure['paragraphs']),
            'section_count': len(structure['sections']),
            'table_count': len(structure['tables'])
        }
