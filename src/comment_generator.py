"""
Comment Generator Module
Inserts contextual comments and suggestions into .docx documents
"""

import logging
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import RGBColor


class CommentGenerator:
    """
    Generates and inserts inline comments into .docx documents.
    All annotations are added as visible text in the document, since python-docx does not support native Word comments.
    """

    def __init__(self) -> None:
        pass

    def add_comment_to_paragraph(self, paragraph, doc: Document, comment_text: str, author: str = "ADGM Agent") -> bool:
        """
        Add an inline annotation to a specific paragraph.
        Returns True if annotation was added successfully.
        """
        try:
            return self._add_inline_annotation(paragraph, comment_text, author)
        except Exception as e:
            logging.error(f"Error adding annotation: {str(e)}")
            return False

    def _add_inline_annotation(self, paragraph, comment_text: str, author: str) -> bool:
        """
        Add inline annotation as appended note in a new paragraph after the target paragraph.
        """
        try:
            comment_para = paragraph._parent.add_paragraph()
            try:
                comment_para.style = 'Comment'
            except Exception:
                try:
                    comment_para.style = 'Normal'
                except Exception:
                    pass
            # Add author and timestamp
            from datetime import date
            timestamp = date.today().isoformat()
            header_run = comment_para.add_run(f"[{author} - {timestamp}]: ")
            header_run.font.color.rgb = RGBColor(128, 128, 128)
            header_run.font.italic = True
            comment_run = comment_para.add_run(comment_text)
            comment_run.font.color.rgb = RGBColor(0, 0, 255)
            separator_para = paragraph._parent.add_paragraph()
            separator_para.add_run("─" * 50)
            try:
                separator_para.style = 'Comment'
            except Exception:
                try:
                    separator_para.style = 'Normal'
                except Exception:
                    pass
            logging.info(f"Annotated paragraph: '{paragraph.text[:40]}...' with comment: '{comment_text[:40]}...'")
            return True
        except Exception as e:
            logging.error(f"Error adding inline annotation: {str(e)}")
            return False
    
    def add_comments_to_document(self, file_path: str, issues: List[Dict[str, Any]], output_path: str) -> bool:
        """
        Add inline comments to document based on identified issues.
        Returns True if document was processed and saved successfully.
        """
        try:
            doc = Document(file_path)
            for issue in issues:
                self._add_issue_comment(doc, issue)
            doc.save(output_path)
            logging.info(f"Annotated document saved to: {output_path}")
            return True
        except Exception as e:
            logging.error(f"Error processing document for comments: {str(e)}")
            return False
    
    def _add_issue_comment(self, doc: Document, issue: Dict[str, Any]) -> None:
        """
        Add an inline comment for a specific issue.
        """
        issue_type = issue.get('type', '')
        severity = issue.get('severity', 'Medium')
        description = issue.get('description', '')
        section = issue.get('section', '')
        adgm_reference = issue.get('adgm_reference', '')
        suggestion = issue.get('suggestion', '')
        
        # Create comment text
        comment_text = self._format_comment_text(
            issue_type, severity, description, 
            section, adgm_reference, suggestion
        )
        
        # Find relevant paragraph to comment on
        target_paragraph = self._find_target_paragraph(doc, issue)
        
        if target_paragraph:
            self.add_comment_to_paragraph(
                target_paragraph, 
                doc,  # Pass the document instance explicitly
                comment_text, 
                "ADGM Corporate Agent"
            )
    
    def _format_comment_text(self, issue_type: str, severity: str, description: str,
                            section: str, adgm_reference: str, suggestion: str) -> str:
        """
        Format comment text with proper structure.
        """
        comment_parts = []
        
        # Issue header
        comment_parts.append(f"ISSUE: {issue_type}")
        comment_parts.append(f"SEVERITY: {severity}")
        comment_parts.append("")
        
        # Description
        comment_parts.append(f"DESCRIPTION: {description}")
        comment_parts.append("")
        
        # Section and reference
        if section:
            comment_parts.append(f"SECTION: {section}")
        if adgm_reference:
            comment_parts.append(f"ADGM REFERENCE: {adgm_reference}")
        comment_parts.append("")
        
        # Suggestion
        if suggestion:
            comment_parts.append(f"SUGGESTION: {suggestion}")
        
        return "\n".join(comment_parts)
    
    def _find_target_paragraph(self, doc: Document, issue: Dict[str, Any]) -> Optional[Any]:
        """
        Find the most relevant paragraph to attach the comment to.
        """
        section = issue.get('section', '')
        clause = issue.get('clause', '')
        
        # Try to find paragraph by section name
        if section:
            for paragraph in doc.paragraphs:
                if section.lower() in paragraph.text.lower():
                    return paragraph
        
        # Try to find paragraph by clause content
        if clause:
            for paragraph in doc.paragraphs:
                if clause.lower() in paragraph.text.lower():
                    return paragraph
        
        # If no specific match, find paragraph with relevant keywords
        keywords = self._extract_keywords_from_issue(issue)
        for paragraph in doc.paragraphs:
            if any(keyword.lower() in paragraph.text.lower() for keyword in keywords):
                return paragraph
        
        # Default to first non-empty paragraph
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                return paragraph
        
        return None
    
    def _extract_keywords_from_issue(self, issue: Dict[str, Any]) -> List[str]:
        """
        Extract relevant keywords from issue for paragraph matching.
        """
        keywords = []
        
        # Extract from description
        description = issue.get('description', '')
        if 'jurisdiction' in description.lower():
            keywords.extend(['jurisdiction', 'courts', 'legal'])
        if 'clause' in description.lower():
            keywords.extend(['clause', 'section', 'article'])
        if 'director' in description.lower():
            keywords.extend(['director', 'board', 'management'])
        if 'shareholder' in description.lower():
            keywords.extend(['shareholder', 'member', 'ownership'])
        
        return keywords
    
    def add_highlight_to_text(self, paragraph, doc: Document, text_to_highlight: str, highlight_color: str = "yellow") -> bool:
        """
        Add highlighting to specific text in a paragraph (not natively supported, so adds an inline annotation).
        """
        try:
            # This is a simplified highlighting approach
            # In a full implementation, you'd need to handle text ranges more carefully
            if text_to_highlight in paragraph.text:
                # For now, we'll just add a comment indicating highlighting is needed
                comment_text = f"HIGHLIGHT: This text should be highlighted: '{text_to_highlight}'"
                self.add_comment_to_paragraph(paragraph, doc, comment_text, "ADGM Agent")
                return True
        except Exception as e:
            print(f"Error adding highlight: {str(e)}")
            return False
        
        return False
    
    def generate_summary_comment(self, doc: Document, compliance_result: Dict[str, Any]) -> bool:
        """
        Add a summary comment at the beginning of the document.
        """
        try:
            if not doc.paragraphs:
                return False
            
            # Create summary comment
            summary_text = self._format_summary_comment(compliance_result)
            
            # Add to first paragraph
            first_paragraph = doc.paragraphs[0]
            self.add_comment_to_paragraph(
                first_paragraph, 
                doc,  # Pass the document instance explicitly
                summary_text, 
                "ADGM Corporate Agent - Summary"
            )
            
            return True
            
        except Exception as e:
            print(f"Error adding summary comment: {str(e)}")
            return False
    
    def _format_summary_comment(self, compliance_result: Dict[str, Any]) -> str:
        """
        Format the summary comment text.
        """
        compliance_score = compliance_result.get('compliance_score', 0)
        total_issues = compliance_result.get('total_issues', 0)
        is_compliant = compliance_result.get('is_compliant', False)
        
        summary_parts = []
        summary_parts.append("DOCUMENT COMPLIANCE SUMMARY")
        summary_parts.append("=" * 30)
        summary_parts.append("")
        
        # Overall status
        if is_compliant:
            summary_parts.append("✅ COMPLIANCE STATUS: COMPLIANT")
        else:
            summary_parts.append("❌ COMPLIANCE STATUS: NON-COMPLIANT")
        
        summary_parts.append("")
        
        # Score and issues
        summary_parts.append(f"COMPLIANCE SCORE: {compliance_score}%")
        summary_parts.append(f"TOTAL ISSUES FOUND: {total_issues}")
        summary_parts.append("")
        
        # Issue breakdown
        high_issues = compliance_result.get('high_severity_issues', 0)
        medium_issues = compliance_result.get('medium_severity_issues', 0)
        low_issues = compliance_result.get('low_severity_issues', 0)
        
        if high_issues > 0:
            summary_parts.append(f"🔴 HIGH SEVERITY: {high_issues}")
        if medium_issues > 0:
            summary_parts.append(f"🟡 MEDIUM SEVERITY: {medium_issues}")
        if low_issues > 0:
            summary_parts.append(f"🟢 LOW SEVERITY: {low_issues}")
        
        summary_parts.append("")
        summary_parts.append("Review all comments below for detailed analysis and suggestions.")
        
        return "\n".join(summary_parts)
